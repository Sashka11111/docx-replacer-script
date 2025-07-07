import os
from docx import Document
from docx.oxml.ns import qn

# 🔷 Шлях до папки з документами
folder_path = r"C:\Users\1\Desktop\КН-41 – копія"

# 🔁 Список замін (текст буде видалено)
replacements = {
    "8.ІНФОРМАЦІЯ ПРО НАЦІОНАЛЬНУ СИСТЕМУ ВИЩОЇ ОСВІТИ": "",
}
# replacements = {
#     "6.2.4 Further information sources": "6.2.5 Further information sources",
# }
# replacements = {
#     "Аccess to study for a Bachelor's Professional pre-higher education educational-professional degree": "Аccess to study for a Bachelor's Professional Degree",
# }
# 🔄 Заміна тексту в параграфі та видалення розриву розділу після нього
def replace_and_clean_after(doc):
    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        paragraph = paragraphs[i]
        text = paragraph.text
        for old, new in replacements.items():
            if old in text:
                # Заміна тексту
                new_text = text.replace(old, new)
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)

                # 🔽 Перевірка наступного параграфа після заміни
                if i + 1 < len(paragraphs):
                    next_p = paragraphs[i + 1]
                    p_elem = next_p._element

                    # Якщо порожній — видаляємо
                    if next_p.text.strip() == "":
                        p_elem.getparent().remove(p_elem)
                        paragraphs.pop(i + 1)
                        continue  # Не інкрементуємо i

                    # Якщо є розрив розділу — видаляємо
                    for child in p_elem.iterchildren():
                        if child.tag == qn('w:pPr'):
                            for sub in list(child):
                                if sub.tag == qn('w:sectPr'):
                                    child.remove(sub)
                                    break
        i += 1

# 🔄 Обробка таблиць (заміна і видалення після)
def process_tables(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_and_clean_after(cell)

# 🔁 Обробка кожного документа
def process_doc(file_path):
    doc = Document(file_path)
    replace_and_clean_after(doc)
    process_tables(doc)
    doc.save(file_path)
    print(f"✅ Оброблено: {os.path.basename(file_path)}")

# 🚀 Обхід усіх файлів
for filename in os.listdir(folder_path):
    if filename.endswith('.docx') and not filename.startswith('~$'):
        process_doc(os.path.join(folder_path, filename))

print("🎉 Готово! Заміни та очищення завершені.")
