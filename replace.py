import os
from docx import Document
from docx.shared import Cm, Pt

# 📁 Шлях до папки
folder_path = r"C:\Users\1\Desktop\ЗВ-41"

# 🔍 Текст для пошуку
target_paragraphs = [
    {
        "ukr": "Назва освітньої кваліфікації та присвоєний освітньо-професійний ступінь (мовою оригіналу)",
        "eng": "Name of educational qualification and educational-professional degree conferred (in original language)",
        "number": "2.1"
    },
    {
        "ukr": "Освітньо-професійний ступінь фахової передвищої освіти",
        "eng": "Professional pre-higher education educational-professional degree",
        "number": "2.1.1"
    }
]

def copy_paragraph_formatting(source_paragraph, target_paragraph, text, indent_cm, bold=True, size_pt=9.5):
    """Копіює форматування абзацу та додає текст із заданим стилем"""
    target_paragraph.paragraph_format.alignment = source_paragraph.paragraph_format.alignment
    target_paragraph.paragraph_format.left_indent = Cm(indent_cm)
    target_paragraph.paragraph_format.right_indent = source_paragraph.paragraph_format.right_indent
    target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
    target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
    target_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing

    target_run = target_paragraph.add_run(text)
    if source_paragraph.runs:
        source_run = source_paragraph.runs[0]
        target_run.font.name = source_run.font.name
        target_run.font.color.rgb = source_run.font.color.rgb
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
    target_run.font.size = Pt(size_pt)
    target_run.font.bold = bold

def replace_exact_lines(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for target in target_paragraphs:
            if target["ukr"] in text and target["eng"] in text:
                # 🧠 Визначаємо, скільки рядків буде в таблиці
                is_211 = target["number"] == "2.1.1"
                table_rows = 2 if is_211 else 1

                # 🔸 Створення таблиці
                table = doc.add_table(rows=table_rows, cols=2)
                table.autofit = False
                section = doc.sections[0]
                page_width = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
                column_width = page_width / 2
                table.columns[0].width = Cm(column_width)
                table.columns[1].width = Cm(column_width)

                # 🔹 Перший рядок — жирний
                cell_ukr_1 = table.cell(0, 0)
                cell_eng_1 = table.cell(0, 1)
                copy_paragraph_formatting(paragraph, cell_ukr_1.paragraphs[0],
                                          f"{target['number']} {target['ukr']}", indent_cm=0.2, bold=True)
                copy_paragraph_formatting(paragraph, cell_eng_1.paragraphs[0],
                                          f"{target['number']} {target['eng']}", indent_cm=0, bold=True)

                # 🔹 Другий рядок — звичайний тільки для 2.1.1
                if is_211:
                    cell_ukr_2 = table.cell(1, 0)
                    cell_eng_2 = table.cell(1, 1)
                    copy_paragraph_formatting(paragraph, cell_ukr_2.paragraphs[0],
                                              "Фаховий молодший бакалавр", indent_cm=0.2, bold=False)
                    copy_paragraph_formatting(paragraph, cell_eng_2.paragraphs[0],
                                              "Professional junior bachelor", indent_cm=0, bold=False)

                # 🗑 Видалити розрив розділу після абзацу, якщо є
                next_el = paragraph._element.getnext()
                if next_el is not None and next_el.tag.endswith('sectPr'):
                    paragraph._element.getparent().remove(next_el)
                    print("🗑 Видалено розрив розділу після абзацу")

                # ➕ Вставити таблицю перед абзацом
                paragraph._element.addprevious(table._element)

                # 🗑 Видалити старий абзац
                paragraph._element.getparent().remove(paragraph._element)

                print(f"✅ Замінено: {target['number']}")
    return doc

def process_docx(file_path):
    try:
        doc = Document(file_path)
        for section in doc.sections:
            print(f"📏 Поля документа: ліве={section.left_margin.cm} см, праве={section.right_margin.cm} см, "
                  f"верхнє={section.top_margin.cm} см, нижнє={section.bottom_margin.cm} см")
        doc = replace_exact_lines(doc)
        output_path = file_path
        doc.save(output_path)
        print(f"💾 Збережено: {os.path.basename(output_path)}")
    except Exception as e:
        print(f"❌ Помилка у {os.path.basename(file_path)}: {str(e)}")

# 🔁 Обробка всіх .docx у папці
for filename in os.listdir(folder_path):
    if filename.endswith(".docx") and not filename.startswith("~$"):
        process_docx(os.path.join(folder_path, filename))

print("🎉 Готово! Таблиці вставлені: пункт 2.1 — один рядок, пункт 2.1.1 — два рядки з кваліфікацією.")
