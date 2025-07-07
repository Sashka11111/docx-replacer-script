import os
from docx import Document
from docx.oxml.ns import qn

# 📁 Шлях до папки з .docx файлами
folder_path = r"C:\Users\1\Desktop\КН-41"

def clean_section_break_and_space(doc):
    for i, paragraph in enumerate(doc.paragraphs):
        if "2.4" in paragraph.text and "Мова" in paragraph.text:
            if i > 0:
                prev_paragraph = doc.paragraphs[i - 1]
                p = prev_paragraph._element
                cleaned = False

                # 🧽 Видалення розриву розділу
                for child in p.iterchildren():
                    if child.tag == qn('w:pPr'):
                        for sub in list(child):
                            if sub.tag == qn('w:sectPr'):
                                child.remove(sub)
                                cleaned = True

                # 🧽 Видалення попереднього параграфа, якщо він порожній або з пробілами
                if prev_paragraph.text.strip() == "":
                    p.getparent().remove(p)
                    cleaned = True

                return cleaned
    return False

# 🔁 Проходимо всі файли в папці
for filename in os.listdir(folder_path):
    if filename.endswith(".docx"):
        file_path = os.path.join(folder_path, filename)
        print(f"📂 Обробка: {filename}")
        try:
            doc = Document(file_path)
            if clean_section_break_and_space(doc):
                doc.save(file_path)  # 💾 Зберігаємо зміни в тому ж файлі
                print("✅ Оновлено:", filename)
            else:
                print("ℹ️ Не знайдено розриву або пробілу для видалення.")
        except Exception as e:
            print(f"❌ Помилка в {filename}: {e}")
