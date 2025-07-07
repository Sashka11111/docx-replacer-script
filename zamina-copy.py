import os
import re
import shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docxcompose.composer import Composer

# 🔷 Шляхи
folder_path = r"C:\Users\1\Desktop\КН-41 – копія"
dyploma_path = r"C:\Users\1\Desktop\test – копія.docx"

replacements = {
   }

keys_to_format = [
    'Офіційна печатка',
    'Official Seal',
]
no_line_break_keys = keys_to_format

def replace_text_in_paragraph(paragraph, doc):
    runs = paragraph.runs
    if not runs:
        return True

    # Зберігаємо оригінальне вирівнювання
    original_alignment = paragraph.alignment

    # Об’єднуємо текст із усіх run-ів
    full_text = ''.join(run.text for run in runs)
    original_text = full_text

    # Виконуємо заміни
    for old, new in replacements.items():
        full_text = full_text.replace(old, new)

    if full_text != original_text:
        print(f"📝 Замінено в параграфі:\n  До: {original_text}\n  Після: {full_text}")

        # Перевіряємо, чи потрібно видалити параграф
        if full_text.strip() == '':
            if not original_text.startswith(("4.3", "4.4")):
                paragraph._element.getparent().remove(paragraph._element)
                return False
        else:
            # Очищаємо існуючі run-и
            for run in runs:
                run.text = ''

            # Додаємо новий run із оновленим текстом
            new_run = paragraph.add_run(full_text)

            # Копіюємо форматування з першого run
            first_run = runs[0]
            new_run.font.name = first_run.font.name
            new_run.font.size = first_run.font.size
            new_run.font.bold = first_run.font.bold
            new_run.font.italic = first_run.font.italic
            if first_run.font.color and first_run.font.color.rgb:
                new_run.font.color.rgb = first_run.font.color.rgb

            # Відновлюємо оригінальне вирівнювання
            paragraph.alignment = original_alignment

            # Встановлюємо JUSTIFY лише для певних ключів
            if any(key in full_text for key in keys_to_format):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return True
    runs = paragraph.runs
    if not runs:
        return True

    # Об'єднуємо текст усіх runs у параграфі
    full_text = ''.join(run.text for run in runs)
    original_text = full_text

    # Виконуємо всі заміни в об'єднаному тексті
    for old, new in replacements.items():
        full_text = full_text.replace(old, new)

    if full_text != original_text:
        print(f"📝 Замінено в параграфі:\n  До: {original_text}\n  Після: {full_text}")

        # Перевіряємо, чи потрібно видалити параграф (якщо новий текст порожній)
        if full_text.strip() == '':
            # Не видаляємо, якщо текст починається з "4.3" або "4.4" (за потреби)
            if not original_text.startswith(("4.3", "4.4")):
                paragraph._element.getparent().remove(paragraph._element)
                return False
        else:
            # Очищаємо всі runs
            for run in runs:
                run.text = ''

            # Додаємо новий run з новим текстом
            new_run = paragraph.add_run(full_text)

            # Копіюємо форматування з першого run
            first_run = runs[0]
            new_run.font.name = first_run.font.name
            new_run.font.size = first_run.font.size
            new_run.font.bold = first_run.font.bold
            new_run.font.italic = first_run.font.italic
            if first_run.font.color and first_run.font.color.rgb:
                new_run.font.color.rgb = first_run.font.color.rgb

            # Застосовуємо вирівнювання, якщо потрібно
            if any(key in full_text for key in keys_to_format):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return True
    runs = paragraph.runs
    if not runs:
        return True

    # Обробка кожного run окремо
    changed = False
    needs_formatting = False
    original_texts = [run.text for run in runs if run.text]  # Зберігаємо оригінальний текст

    for run in runs:
        run_text = run.text
        if not run_text:
            continue
        for old, new in replacements.items():
            if old in run_text:
                run_text = run_text.replace(old, new)
                changed = True
                if old in keys_to_format or new in keys_to_format:
                    needs_formatting = True
        run.text = run_text

    # Перевірка на видалення параграфа
    full_text = ''.join(run.text for run in runs).strip()
    for old, new in replacements.items():
        if new == '' and old == full_text:  # Точний збіг усього тексту
            if full_text.startswith("4.3") or full_text.startswith("4.4"):
                continue
            paragraph._element.getparent().remove(paragraph._element)
            return False

    if changed:
        print(f"📝 Замінено в параграфі:\n  До: {''.join(original_texts)}\n  Після: {full_text}")

        # Зберігаємо стилі з першого run
        first_run = runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
        font_bold = first_run.font.bold
        font_italic = first_run.font.italic
        font_color = first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None

        # Застосовуємо стилі до всіх run
        for run in runs:
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = font_bold
            run.font.italic = font_italic
            if font_color:
                run.font.color.rgb = font_color

        # Застосовуємо вирівнювання, якщо потрібно
        if needs_formatting and full_text not in no_line_break_keys:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return True

def remove_page_break_before_625(doc):
    pattern = re.compile(r'\b6\.2\.5\b')
    for i in range(1, len(doc.paragraphs)):
        paragraph = doc.paragraphs[i]
        if pattern.search(paragraph.text):
            prev_para = doc.paragraphs[i - 1]
            # Якщо в попередньому абзаці є тільки розрив сторінки — видалити його
            if any('pageBreakBefore' in run._element.xml for run in prev_para.runs) or prev_para.text.strip() == '':
                print(f"✂️ Видалено розрив перед '6.2.5'")
                prev_para._element.getparent().remove(prev_para._element)
            break


def replace_text_in_cell(cell, doc):
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, doc)

def delete_after_section(doc):
    pattern = re.compile(r'ІНФОРМАЦІЯ\s*ПРО\s*НАЦІОНАЛЬНУ\s*СИСТЕМУ\s*ВИЩОЇ', re.IGNORECASE | re.UNICODE)
    start_index = None
    for i, paragraph in enumerate(doc.paragraphs):
        full_text = ''.join(run.text for run in paragraph.runs).strip()
        if pattern.search(full_text):
            start_index = i
            break

    if start_index is not None:
        try:
            for i in reversed(range(start_index + 1, len(doc.paragraphs))):
                doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
            for table in reversed(doc.tables):
                if table._element.getparent().index(table._element) > start_index:
                    table._element.getparent().remove(table._element)
            print("🗑️ Видалено розділи після 'ІНФОРМАЦІЯ ПРО НАЦІОНАЛЬНУ СИСТЕМУ ВИЩОЇ ОСВІТИ'")
        except Exception as e:
            print(f"❌ Помилка при видаленні розділу: {str(e)}")

def append_dyploma_with_formatting(original_path, dyploma_path):
    try:
        if not os.path.exists(dyploma_path):
            print(f"❌ Файл dyploma.docx не знайдено за шляхом {dyploma_path}")
            return

        temp_path = original_path.replace(".docx", "_temp.docx")
        shutil.copyfile(original_path, temp_path)
        try:
            base_doc = Document(temp_path)
            composer = Composer(base_doc)
            dyploma_doc = Document(dyploma_path)
            composer.append(dyploma_doc)
            composer.save(original_path)
            print(f"📎 Вставлено dyploma.docx у {os.path.basename(original_path)}")
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
    except Exception as e:
        print(f"❌ Помилка вставки dyploma.docx у {os.path.basename(original_path)}: {str(e)}")

def process_docx(file_path, dyploma_path):
    try:
        if not os.access(file_path, os.R_OK | os.W_OK):
            print(f"❌ Немає прав доступу до {os.path.basename(file_path)}")
            return

        doc = Document(file_path)
    except Exception as e:
        print(f"❌ Помилка при відкритті {os.path.basename(file_path)}: {str(e)}")
        return

    for paragraph in doc.paragraphs:
        if not replace_text_in_paragraph(paragraph, doc):
            continue

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_cell(cell, doc)
    remove_page_break_before_625(doc)
    delete_after_section(doc)

    try:
        doc.save(file_path)
        print(f"💾 Збережено: {os.path.basename(file_path)}")
    except Exception as e:
        print(f"❌ Помилка при збереженні {os.path.basename(file_path)}: {str(e)}")
        return

    append_dyploma_with_formatting(file_path, dyploma_path)

    print(f"✅ Оброблено: {os.path.basename(file_path)}")

try:
    if not os.path.exists(folder_path):
        print(f"❌ Папка {folder_path} не існує")
    else:
        for filename in os.listdir(folder_path):
            if filename.endswith('.docx') and not filename.startswith('~$'):
                full_path = os.path.join(folder_path, filename)
                print(f"\n📄 Починаємо: {filename}")
                process_docx(full_path, dyploma_path)
        print("\n🎉 Обробка всіх файлів завершена!")
except Exception as e:
    print(f"❌ Помилка при доступі до папки {folder_path}: {str(e)}")
