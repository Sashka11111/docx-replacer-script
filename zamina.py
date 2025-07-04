import os
import re
import shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docxcompose.composer import Composer

# 🔷 Шляхи
folder_path = r"C:\Users\1\Desktop\Навчальна частина\ПРз-31"
dyploma_path = r"C:\Users\1\Desktop\test.docx"

# 🔁 Словник замін
replacements = {
    'Назва кваліфікації та присвоєний ступінь': 'Назва освітньої кваліфікації та присвоєний освітньо-професійний ступінь (мовою оригіналу)',
    'Name of qualification and (if applicable) title conferred': 'Name of educational qualification and educational-professional degree conferred (in original language)',
    'Ступінь вищої освіти': 'Освітньо-професійний ступінь фахової передвищої освіти',
    'Degree': 'Professional pre-higher education educational-professional degree',
    'Професійна кваліфікація (у разі присвоєння)': 'Освітньо-професійна програма',
    'Professional Qualification (if awarded)': 'Educational-professional programme',

#nen
    '–': 'Правознавство',

    'Основна (основні) галузь (галузі) знань за кваліфікацією': 'Професійна кваліфікація (у разі присвоєння)',
    'Main field(s) of study for the qualification': 'Professional qualification (if awarded)',

#nen
    '08 Право': '–',
    '08 Law': '–',


    'Найменування і статус закладу (якщо відмінні від п. 2.3), який реалізує освітню програму': '',
    'Name and status of institution (if different from 2.3)': '',
    'administering studies': '',
    'Зазначено у пункті 2.3': '',
    'Specified in 2.3': '',
    '2.4': '',
    '2.5': '2.4',
    '6.2.5': '6.2.5',
    'ПРО РІВЕНЬ КВАЛІФІКАЦІЇ': 'ПРО КВАЛІФІКАЦІЮ',
    'INFORMATION ON THE LEVEL AND DURATION OF THE QUALIFICATION': 'INFORMATION ON THE LEVEL OF QUALIFICATION AND LENGTH OF PROGRAMME',
    'Тривалість освітньої програми в кредитах та/або роках': 'Офіційна тривалість освітньо-професійної програми в кредитах та/або роках',
    'Official duration of programme in credits and/or years': 'Official length of educational-professional programme in credits and/or years',
    'ІНФОРМАЦІЯ ПРО ЗАВЕРШЕНУ ОСВІТНЮ ПРОГРАМУ ТА ЗДОБУТІ РЕЗУЛЬТАТИ НАВЧАННЯ': 'ІНФОРМАЦІЯ ПРО ЗАВЕРШЕНУ ОСВІТНЬО-ПРОФЕСІЙНУ ПРОГРАМУ ТА ЗДОБУТІ РЕЗУЛЬТАТИ НАВЧАННЯ',
    'INFORMATION ON THE PROGRAMME COMPLETED AND THE RESULTS OBTAINED': 'INFORMATION ON THE COMPLETED EDUCATIONAL-PROFESSIONAL PROGRAMME AND LEARNING OUTCOMES',
    'Найменування всіх закладів вищої освіти (наукових установ) (відокремлених структурних підрозділів закладів вищої освіти), у яких здобувалася кваліфікація (у тому числі заклади освіти, в яких здобувач вищої освіти вивчав окремі дисципліни за програмами академічної мобільності)': 'Найменування всіх закладів фахової передвищої освіти (структурних підрозділів або філій закладів фахової передвищої освіти), у яких здобувалася освітня кваліфікація (у тому числі заклади освіти, в яких здобувач фахової передвищої освіти вивчав окремі дисципліни за програмами академічної мобільності)',
  'Name of all higher education (research)  s (separate structural units of higher education s) where the qualification has been gained (including education s where the holder of the qualification has been studying separate course units within the framework(s) of academic mobility)':
  'Names of all higher education (research) institutions (separate structural units of higher education institutions) where the qualification has been gained (including education institutions where the holder of the qualification has been studying separate course units within the framework(s) of academic mobility)',


    'закладу вищої освіти (наукової установи)': 'закладу фахової передвищої освіти (іншого суб’єкта освітньої діяльності)',
    'institution':'',
    'Contact information of the higher education (research)':'Contact information of the professional pre-higher education institution (other educational entity)',
    'Керівник або уповноважена особа закладу вищої освіти': 'Керівник або уповноважена особа закладу фахової передвищої освіти',
    'Capacity': 'Head or other authorized person of professional pre-higher education institution',
    'Посада керівника або іншої уповноваженої особи закладу вищої освіти (наукової установи)': 'Посада керівника або іншої уповноваженої особи закладу фахової передвищої освіти',
    'Position of the Head or another authorized person of the Higher Education (Research) Institution': 'Position of the professional pre-higher education institution head or other authorized person',
    'Печатка': 'Офіційна печатка',
    'Official stamp or seal': 'Official Seal',
    '8. ІНФОРМАЦІЯ ПРО НАЦІОНАЛЬНУ СИСТЕМУ ВИЩОЇ ОСВІТИ': '',
    'ІНФОРМАЦІЯ ПРО СИСТЕМУ ФАХОВОЇ ПЕРЕДВИЩОЇ ОСВІТИ': '',
}

keys_to_format = ['Офіційна печатка', 'Official Seal']

def replace_text_in_cell(cell):
    for paragraph in cell.paragraphs:
        runs = paragraph.runs
        if not runs:
            continue

        full_text_before = ''.join(run.text for run in runs)
        full_text_after = full_text_before

        for old, new in replacements.items():
            full_text_after = full_text_after.replace(old, new)

        if full_text_after != full_text_before:
            print(f"🔁 Заміна в таблиці:\n  До: {full_text_before}\n  Після: {full_text_after}")

            first_run = runs[0]
            for run in runs:
                run.text = ''

            new_run = paragraph.add_run(full_text_after)
            new_run.font.name = first_run.font.name
            new_run.font.size = first_run.font.size
            new_run.font.bold = first_run.font.bold
            new_run.font.italic = first_run.font.italic
            if first_run.font.color and first_run.font.color.rgb:
                new_run.font.color.rgb = first_run.font.color.rgb

            if any(key in full_text_after for key in keys_to_format):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def process_docx(file_path, dyploma_path):
    try:
        doc = Document(file_path)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_text_in_cell(cell)

        doc.save(file_path)
        print(f"💾 Збережено: {os.path.basename(file_path)}")
    except Exception as e:
        print(f"❌ Помилка в {file_path}: {e}")

def main():
    if not os.path.exists(folder_path):
        print(f"❌ Папка {folder_path} не існує")
        return

    for filename in os.listdir(folder_path):
        if filename.endswith('.docx') and not filename.startswith('~$'):
            full_path = os.path.join(folder_path, filename)
            print(f"\n📄 Починаємо: {filename}")
            process_docx(full_path, dyploma_path)

    print("\n🎉 Обробка завершена")

if __name__ == '__main__':
    main()
